from operator import itemgetter
from itertools import cycle
from docx import Document
import warnings
import fitz
import json
import re


def fonts(doc, granularity=False):
    """Extracts fonts and their usage in PDF documents.
    :param doc: PDF document to iterate through
    :type doc: <class 'fitz.fitz.Document'>
    :param granularity: also use 'font', 'flags' and 'color' to discriminate text
    :type granularity: bool
    :rtype: [(font_size, count), (font_size, count}], dict
    :return: most used fonts sorted by count, font style information
    """
    styles = {}
    font_counts = {}
    page_number = 0

    for page in doc:
        page_number = page_number + 1
        blocks = page.getText("dict")["blocks"]
        for b in blocks:  # iterate through the text blocks
            if b['type'] == 0:  # block contains text
                for l in b["lines"]:  # iterate through the text lines
                    for s in l["spans"]:  # iterate through the text spans
                        if granularity:
                            identifier = "{0}_{1}_{2}_{3}".format(s['size'], s['flags'], s['font'], s['color'])
                            styles[identifier] = {'size': s['size'], 'flags': s['flags'], 'font': s['font'],
                                               'color': s['color']}
                        else:
                            identifier = "{0}".format(s['size'])
                            styles[identifier] = {'size': s['size'], 'font': s['font']}

                        font_counts[identifier] = font_counts.get(identifier, 0) + 1  # count the fonts usage

    font_counts = sorted(font_counts.items(), key=itemgetter(1), reverse=True)

    if len(font_counts) < 1:
        raise ValueError("Zero discriminating fonts found!")

    return font_counts, styles


def font_tags(font_counts, styles):
    """Returns dictionary with font sizes as keys and tags as value.
    :param font_counts: (font_size, count) for all fonts occuring in document
    :type font_counts: list
    :param styles: all styles found in the document
    :type styles: dict
    :rtype: dict
    :return: all element tags based on font-sizes
    """
    p_style = styles[font_counts[0][0]]  # get style for most used font by count (paragraph)
    p_size = p_style['size']  # get the paragraph's size

    # sorting the font sizes high to low, so that we can append the right integer to each tag
    font_sizes = []
    for (font_size, count) in font_counts:
        font_sizes.append(float(font_size))
    font_sizes.sort(reverse=True)

    # aggregating the tags for each font size
    idx = 0
    size_tag = {}
    for size in font_sizes:
        idx += 1
        if size == p_size:
            idx = 0
            size_tag[size] = '<p>'
        if size > p_size:
            size_tag[size] = '<HEADER_{0}>'.format(idx)
        elif size < p_size:
            size_tag[size] = '<s{0}>'.format(idx)

    return size_tag


def headers_para(doc, size_tag,starting, ending):
    """Scrapes headers & paragraphs from PDF and return texts with element tags.
    :param doc: PDF document to iterate through
    :type doc: <class 'fitz.fitz.Document'>
    :param size_tag: textual element tags for each size
    :type size_tag: dict
    :rtype: list
    :return: texts with pre-pended element tags
    """
    header_para = []  # list with headers and paragraphs
    first = True  # boolean operator for first header
    previous_s = {}  # previous span

    for page in doc.pages(starting,ending,1): #64
        blocks = page.getText("dict")["blocks"]
        for b in blocks:  # iterate through the text blocks
            if b['type'] == 0:  # this block contains text

                # REMEMBER: multiple fonts and sizes are possible IN one block

                block_string = ""  # text found in block
                for l in b["lines"]:  # iterate through the text lines
                    for s in l["spans"]:  # iterate through the text spans
                        if s['text'].strip():  # removing whitespaces:
                            if first:
                                previous_s = s
                                first = False
                                block_string = size_tag[s['size']] + s['text']
                            else:
                                if s['size'] == previous_s['size']:

                                    if block_string and all((c == "|") for c in block_string):
                                        # block_string only contains pipes
                                        block_string = size_tag[s['size']] + s['text']
                                    if block_string == "":
                                        # new block has started, so append size tag
                                        block_string = size_tag[s['size']] + s['text']
                                    else:  # in the same block, so concatenate strings
                                        block_string += " " + s['text']

                                else:
                                    header_para.append(block_string)
                                    block_string = size_tag[s['size']] + s['text']

                                previous_s = s

                    # new block started, indicating with a pipe
                    block_string += ""
                header_para.append(block_string)

    return header_para


def is_ascii(s):
    """Returns char if it is ascii.
    :param s: single character of a string
    :type : char
    :rtype: char
    :return: ascii character
    """
    return all(ord(c) < 128 for c in s)


print("Enter file name without extension.")
doc_name = input('Enter your pdf file name here: ')
print("Starting page number has to be greater than 0.")
start_str = input('Enter starting page number: ')
end_str = input('Enter ending page number: ')

document = doc_name + '.pdf'
starting = int(start_str)
ending = int(end_str)
elements = []

doc = fitz.open(document)
document_save = Document()
font_counts, styles = fonts(doc, granularity=False)
size_tag = font_tags(font_counts, styles)
elements = headers_para(doc, size_tag, starting - 1, ending)

print("Total lines are " + str(len(elements)))
print("Adding the header font tags from page "+start_str+" to "+end_str)
document_save.add_paragraph("List of the header font tags from page "+start_str+" to "+end_str+" with headers")

for each_element in elements:
    if "<HEADER_" in each_element:
        try:
            document_save.add_paragraph(each_element)
            document_save.save("tag.docx")
        except Exception as e:
            new_str = ""
            print(e)
            print(each_element)
            try:
                for chars in each_element:
                    if is_ascii(chars):
                        new_str = new_str + chars
                document_save.add_paragraph(new_str)
                document_save.save("tag.docx")
            except Exception as en:
                print(en)

print("\n\nPlease check tag.docx file to get the header tags for different header fonts.\n")